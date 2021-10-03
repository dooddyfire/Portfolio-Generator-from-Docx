#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Aug  5 23:22:48 2020

@author: sweeties01
"""
import docx
from docx import Document 
from docx.shared import RGBColor ,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os
import sys
from tkinter import *
from tkinter import messagebox
from tkinter import filedialog
import webbrowser
def clear(e): 
    ent.delete(0,END)
    ent1.delete(0,END)
    ent2.delete(0,END)
    lb['text']=""
def cleared():
    ent.delete(0,END)
    ent1.delete(0,END)
    ent2.delete(0,END)
    ent3.delete(0,END)
    lb['text']=""

# ฟังก์ชันสร้าง portfolio.docx เป็นไฟล์ words (ก็อปโค้ดฟังก์ชัน create_port พร้อมไฟล์ word ที่ gen ออกมาจากโปรแกรม แล้วส่งได้เลย)
def create_port():
    doc = Document()
    
    para = doc.add_heading('PortFolio',level=0)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    try:
        img = doc.add_picture(im,width=docx.shared.Inches(2),height = docx.shared.Inches(2))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass
        
    para1 = doc.add_heading('Name : ',level=2)
    runp1 = para1.add_run(name.get())
    runp1.font.color.rgb = RGBColor(0,0,255)
    runp1.style = 'Emphasis'
    runp1.underline = True 
    runp1.bold = True
    
    para2run = doc.add_paragraph().add_run("Age : {}".format(age.get()))
    print(type(para2run))
    para2run.font.color.rgb = RGBColor(0,0,255)
    para2run.underline = True 
    para2run.bold = True 
    
    para3 = doc.add_paragraph()
    para3.alingment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = para3.add_run("Experience : {}".format(exp.get()))
    run3.font.color.rgb = RGBColor(0,0,255) #เปลี่ยนสี font
    run3.bold = True
    run3.underline =True
    
    para4 = doc.add_heading(f'Description : {des.get()}',level=2)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    
    para5 = doc.add_heading(f'Thank You For Reading',level=0) #เพิ่ม heading level = 0 มันจะมีขัดเส้นใต้
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER #จัดกลาง
    
    doc.add_page_break() # เริ่มหน้าใหม่
    para4 = doc.add_heading('New Page',level=0)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    para5= doc.add_paragraph().add_run('Test Test New Page')
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para5.font.color.rgb = RGBColor(0,0,255)
    para5.font.name = 'Calibri'
    para5.font.size = Pt(30)
    
    para5.bold =True 
    para5.underline = True

    
    doc.save('PortFolio.docx') #Save ลงไฟล์ word ชื่อ PortFolio
    messagebox.showinfo(title="สร้างไฟล์ words สำเร็จ",message="สร้างไฟล์ words สำเร็จ")
    webbrowser.open("file:///" + os.getcwd() + "/PortFolio.docx") #สั่งให้เปิดไฟล์ word
    
def load_im():
    global im
    im = filedialog.askopenfilename()
    lb['text'] = im

def shift2(): 
    x1,y1,x2,y2 = canvas2.bbox('marquee')
    if c<0 or b<0: 
        x1 = canvas2.winfo_width()
        y1 = canvas2.winfo_height()//2 
        canvas2.coords('marquee',x1,y1)
    else: 
        canvas2.move('marquee',-2,0)
    canvas2.after(1000//f,shift2)
    
def shift(): 
    x1,y1,x2,y2 = canvas.bbox('uno')
    if x2<0 or y1<0: 
        x1 = canvas.winfo_width()
        y1 = canvas.winfo_height()//2 
        canvas.coords('uno',x1,y1)
    else: 
        canvas.move('uno',-2,0)
    canvas.after(1000//fps,shift)

def close(): 
    pass
    
if __name__ == '__main__':
    # nam = input('Enter your name : ')
    # ag = input('Enter your age : ')
    # exp = input('exp : ')
    root = Tk()
    root.title('PortFolio Generator')
    
    name = StringVar()
    age = StringVar()
    exp = StringVar()
    im = StringVar()
    des = StringVar()
    
    canvas = Canvas(root,bg='black')
    canvas.pack(fill=X)
    
    textvar = "กรอกข้อมูลตามช่องข้างล่างแล้วกด generate จะได้ไฟล์ words ตามต้องการ"
    text = canvas.create_text(-2000,0,text=textvar,fill='yellow',tags='uno',anchor='w',font=('Arial',15,'bold'))
    x1,y1,x2,y2 = canvas.bbox('uno')
    canvas['width'] = x2-x1
    canvas['height'] = y2-y1 
    fps=50 
    shift()
    
    frame = Frame(root,bg='powder blue',relief='ridge',bd=3)
    frame.pack()
    Label(frame,text="Name",font=('Arial',20,'bold'),fg='blue',bg='powder blue').grid(row=0,column=0,padx=10)
    ent = Entry(frame,textvariable=name,justify=LEFT,width=50)
    ent.grid(row=0,column=1,columnspan=2)
    
    Label(frame,text="AGE",font=('Arial',20,'bold'),fg='red',bg='powder blue').grid(row=1,column=0,padx=10)
    ent1 = Entry(frame,textvariable=age,justify=LEFT,width=50)
    ent1.grid(row=1,column=1,columnspan=2)
    
    Label(frame,text="Experience",font=('Arial',20,'bold'),fg='green',bg='powder blue').grid(row=2,column=0,padx=10)
    ent2 = Entry(frame,textvariable=exp,justify=LEFT,width=50)
    ent2.grid(row=2,column=1,columnspan=2)
    
    Label(frame,text="Description",font=('Arial',20,'bold'),fg='black',bg='powder blue').grid(row=3,column=0,padx=10)
    ent3 = Entry(frame,textvariable=des,justify=LEFT,width=50)
    ent3.grid(row=3,column=1,columnspan=2)
    
    Label(frame,text="image path",font=('Arial',20,'bold'),fg='green',bg='powder blue').grid(row=4,column=0,padx=10)
    lb = Label(frame,text="",justify=LEFT,width=50)
    lb.grid(row=4,column=1,columnspan=2)
    
    btn2 = Button(frame,text="Load Image",command=load_im,cursor='hand2',bd=3)
    btn2.grid(row=5,column=2)
    
    btn = Button(frame,text="Generate",command=create_port,cursor='hand2',bd=3)
    btn.grid(row=5,column=1)
    
    menubar = Menu(root)
    filemenu = Menu(menubar,tearoff=0)
    menubar.add_cascade(label='File',menu=filemenu)
    filemenu.add_command(label='Clear',command=cleared)
    filemenu.add_command(label='Close',command=close)
    root.config(menu=menubar)
    

    root.mainloop()
    # create_port(nam,ag,exp)